from flask import Flask, render_template, request, send_file, redirect, url_for
from PIL import Image
import os, zipfile, shutil, subprocess
import pikepdf
from docx import Document
from openpyxl import Workbook
import pdfplumber
import fitz  # PyMuPDF
from reportlab.pdfgen import canvas
import pandas as pd

app = Flask(__name__)
app.secret_key = 'your_secret_key_here'
UPLOAD_FOLDER = "uploads"
OUTPUT_FOLDER = "output"
TEMP_FOLDER = "temp_extract"

os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(OUTPUT_FOLDER, exist_ok=True)
os.makedirs(TEMP_FOLDER, exist_ok=True)

# Supported formats
SUPPORTED_REDUCE = {'png', 'jpg', 'jpeg', 'bmp', 'pdf', 'docx', 'xlsx'}

# ========== COMPRESSION FUNCTIONS ==========
def compress_image(input_path, output_path, quality=70):
    img = Image.open(input_path)
    img = img.convert("RGB")
    img.save(output_path, optimize=True, quality=quality)

def compress_pdf(input_path, output_path):
    try:
        subprocess.run([
            "gs", "-sDEVICE=pdfwrite",
            "-dCompatibilityLevel=1.4",
            "-dPDFSETTINGS=/ebook",
            "-dNOPAUSE", "-dQUIET", "-dBATCH",
            f"-sOutputFile={output_path}", input_path
        ], check=True)
    except FileNotFoundError:
        pdf = pikepdf.open(input_path)
        pdf.save(output_path)

def compress_office_file(input_path, output_path, quality=70):
    if os.path.exists(TEMP_FOLDER):
        shutil.rmtree(TEMP_FOLDER)
    os.makedirs(TEMP_FOLDER)

    with zipfile.ZipFile(input_path, 'r') as zip_ref:
        zip_ref.extractall(TEMP_FOLDER)

    media_path = os.path.join(TEMP_FOLDER, 'word', 'media')
    if not os.path.exists(media_path):
        media_path = os.path.join(TEMP_FOLDER, 'xl', 'media')

    if os.path.exists(media_path):
        for img_name in os.listdir(media_path):
            img_path = os.path.join(media_path, img_name)
            try:
                img = Image.open(img_path)
                img = img.convert("RGB")
                img.save(img_path, optimize=True, quality=quality)
            except:
                pass

    with zipfile.ZipFile(output_path, 'w', zipfile.ZIP_DEFLATED) as zip_out:
        for root, _, files in os.walk(TEMP_FOLDER):
            for file in files:
                filepath = os.path.join(root, file)
                arcname = os.path.relpath(filepath, TEMP_FOLDER)
                zip_out.write(filepath, arcname)

    shutil.rmtree(TEMP_FOLDER)

# ========== CONVERSION FUNCTIONS ==========
def pdf_to_word(pdf_path, output_path):
    doc = Document()
    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                doc.add_paragraph(text)
    doc.save(output_path)

def word_to_pdf(docx_path, output_path):
    doc = Document(docx_path)
    text = "\n".join([para.text for para in doc.paragraphs])
    
    c = canvas.Canvas(output_path)
    text_obj = c.beginText(40, 800)
    text_obj.setFont("Helvetica", 12)
    
    for line in text.split('\n'):
        text_obj.textLine(line)
    
    c.drawText(text_obj)
    c.save()

def pdf_to_excel(pdf_path, output_path):
    wb = Workbook()
    ws = wb.active
    with pdfplumber.open(pdf_path) as pdf:
        for i, page in enumerate(pdf.pages):
            text = page.extract_text()
            if text:
                for line in text.split('\n'):
                    ws.append([line])
    wb.save(output_path)

def excel_to_pdf(xlsx_path, output_path):
    df = pd.read_excel(xlsx_path)
    c = canvas.Canvas(output_path)
    text_obj = c.beginText(40, 800)
    text_obj.setFont("Helvetica", 10)
    
    text_obj.textLine("\t".join(df.columns))
    
    for _, row in df.iterrows():
        text_obj.textLine("\t".join([str(x) for x in row]))
    
    c.drawText(text_obj)
    c.save()

def image_to_pdf(image_path, output_path):
    image = Image.open(image_path)
    image.convert('RGB').save(output_path)

def pdf_to_image(pdf_path, output_path):
    doc = fitz.open(pdf_path)
    page = doc.load_page(0)
    pix = page.get_pixmap()
    img_format = output_path.split('.')[-1].upper()
    if img_format not in ['PNG', 'JPG', 'JPEG']:
        img_format = 'PNG'
    pix.save(output_path, img_format)

# ========== ROUTES ==========
@app.route("/")
def index():
    message = request.args.get('message')
    message_type = request.args.get('type')
    
    if message and message_type:
        return render_template("index.html", message={
            'text': message,
            'type': message_type
        })
    return render_template("index.html")

@app.route("/upload", methods=["POST"])
def upload_file():
    file = request.files.get("file")

    if not file or file.filename == '':
        return redirect(url_for('index', message="No file selected", type="error"))

    file_ext = file.filename.lower().split(".")[-1]
    if file_ext not in SUPPORTED_REDUCE:
        return redirect(url_for('index', 
            message="This file type is not supported for free users. Contact Developer Andee", 
            type="error"))

    try:
        filepath = os.path.join(UPLOAD_FOLDER, file.filename)
        file.save(filepath)
        output_path = os.path.join(OUTPUT_FOLDER, f"compressed_{file.filename}")

        if file_ext in ["png", "jpg", "jpeg", "bmp"]:
            compress_image(filepath, output_path)
        elif file_ext == "pdf":
            compress_pdf(filepath, output_path)
        elif file_ext in ["docx", "xlsx"]:
            compress_office_file(filepath, output_path)

        return send_file(output_path, as_attachment=True)
    
    except Exception as e:
        return redirect(url_for('index', 
            message=f"Error processing file: {str(e)}", 
            type="error"))

@app.route("/convert", methods=["POST"])
def convert_file():
    file = request.files.get("file")
    conversion_type = request.form.get("target_format")

    if not file or file.filename == '':
        return redirect(url_for('index', message="No file selected", type="error"))
    
    if not conversion_type:
        return redirect(url_for('index', message="No conversion type selected", type="error"))

    file_ext = file.filename.lower().split(".")[-1]
    
    try:
        filepath = os.path.join(UPLOAD_FOLDER, file.filename)
        file.save(filepath)
        
        base_name = os.path.splitext(file.filename)[0]
        output_ext = conversion_type.split('_to_')[-1]
        output_ext = 'jpg' if output_ext == 'img' else output_ext
        output_filename = f"converted_{base_name}.{output_ext}"
        output_path = os.path.join(OUTPUT_FOLDER, output_filename)

        if conversion_type == 'pdf_to_docx':
            if file_ext != 'pdf':
                raise ValueError("Please upload a PDF file for PDF to Word conversion")
            pdf_to_word(filepath, output_path)
        
        elif conversion_type == 'docx_to_pdf':
            if file_ext != 'docx':
                raise ValueError("Please upload a Word (DOCX) file for Word to PDF conversion")
            word_to_pdf(filepath, output_path)
        
        elif conversion_type == 'pdf_to_xlsx':
            if file_ext != 'pdf':
                raise ValueError("Please upload a PDF file for PDF to Excel conversion")
            pdf_to_excel(filepath, output_path)
        
        elif conversion_type == 'xlsx_to_pdf':
            if file_ext != 'xlsx':
                raise ValueError("Please upload an Excel (XLSX) file for Excel to PDF conversion")
            excel_to_pdf(filepath, output_path)
        
        elif conversion_type == 'img_to_pdf':
            if file_ext not in ['jpg', 'jpeg', 'png']:
                raise ValueError("Please upload an image (JPG/JPEG/PNG) for Image to PDF conversion")
            image_to_pdf(filepath, output_path)
        
        elif conversion_type == 'pdf_to_img':
            if file_ext != 'pdf':
                raise ValueError("Please upload a PDF file for PDF to Image conversion")
            output_path = os.path.join(OUTPUT_FOLDER, f"converted_{base_name}.png")
            pdf_to_image(filepath, output_path)
        
        elif conversion_type == 'jpg_to_png':
            if file_ext not in ['jpg', 'jpeg']:
                raise ValueError("Please upload a JPG/JPEG file for JPG to PNG conversion")
            img = Image.open(filepath)
            img.save(output_path, 'PNG')
        
        elif conversion_type == 'png_to_jpg':
            if file_ext != 'png':
                raise ValueError("Please upload a PNG file for PNG to JPG conversion")
            img = Image.open(filepath)
            img.convert('RGB').save(output_path, 'JPEG')
        
        else:
            raise ValueError("Selected conversion is not supported")

        return send_file(output_path, as_attachment=True)
    
    except Exception as e:
        return redirect(url_for('index', 
            message=f"Error during conversion: {str(e)}", 
            type="error"))

if __name__ == "__main__":
    app.run(debug=True)